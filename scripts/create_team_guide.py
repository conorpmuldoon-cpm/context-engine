"""Generate the Context Engine Team Guide as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent.parent / "outputs" / "Context_Engine_Team_Guide.docx"


def set_style(doc):
    """Configure base document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        h.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        h.paragraph_format.space_after = Pt(4)
        if level == 1:
            h.font.size = Pt(16)
        elif level == 2:
            h.font.size = Pt(13)
        else:
            h.font.size = Pt(11)


def add_bullet(doc, text, bold_prefix=None, indent=0.25):
    """Add a bullet-point paragraph, optionally with a bold lead-in."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10.5)
        run_b.font.name = "Calibri"
        p.add_run(text).font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"


def build():
    doc = Document()
    set_style(doc)

    # -- Narrow margins for two-pager --
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ================================================================
    # TITLE
    # ================================================================
    title = doc.add_heading("Context Engine", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Team Guide — Office of Analytics, Performance & Innovation")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "Calibri"
    subtitle.paragraph_format.space_after = Pt(8)

    # ================================================================
    # WHAT IS IT
    # ================================================================
    doc.add_heading("What Is the Context Engine?", level=2)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "The Context Engine is a persistent intelligence tool that collects, organizes, "
        "and synthesizes public-source information relevant to City of Syracuse government operations. "
        "It runs continuously in the background — scanning council meetings, local news, city press releases, "
        "and the syr.gov website — so the Innovation Team always has current, structured context available "
        "when preparing for engagements, briefing leadership, or tracking emerging issues."
    ).font.size = Pt(10.5)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2.add_run(
        "As of March 2026, the Context Engine contains approximately "
    ).font.size = Pt(10.5)
    run_b = p2.add_run("690 records")
    run_b.bold = True
    run_b.font.size = Pt(10.5)
    p2.add_run(
        " spanning council meeting segments, news articles, press releases, and public notices — "
        "organized into "
    ).font.size = Pt(10.5)
    run_b2 = p2.add_run("59 thematic clusters")
    run_b2.bold = True
    run_b2.font.size = Pt(10.5)
    p2.add_run(
        " (e.g., Budget Battle, Lead Remediation, SHA Governance, School Zone Cameras). "
        "Everything comes from publicly accessible sources."
    ).font.size = Pt(10.5)

    # ================================================================
    # HOW DOES IT WORK
    # ================================================================
    doc.add_heading("How Does It Work?", level=2)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(4)
    p3.add_run(
        "Four automated collector agents run every morning and feed information into the system:"
    ).font.size = Pt(10.5)

    # Collector table
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    headers = ["Collector", "What It Does", "Schedule"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)

    rows_data = [
        ("News Scanner", "Scans Syracuse.com RSS for local government news", "Daily"),
        ("Website Monitor", "Checks syr.gov for new press releases and public notices", "Daily"),
        ("Council Transcriber", "Detects new council meeting videos, pulls transcripts, segments by topic", "Daily"),
        ("Agenda Scanner", "Downloads and parses council/committee meeting agendas", "Daily"),
    ]
    for i, (col1, col2, col3) in enumerate(rows_data):
        row = table.rows[i + 1]
        row.cells[0].text = col1
        row.cells[1].text = col2
        row.cells[2].text = col3
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(4)
    p4.add_run("Each piece of information goes through a standard pipeline:").font.size = Pt(10.5)

    steps = [
        ("Collect: ", "The raw article, transcript segment, or press release is captured."),
        ("Enrich: ", "AI (Claude) reads the content and generates a summary, identifies departments, "
         "people, organizations, topic tags, sentiment, and political signals."),
        ("Tag & Cross-Reference: ", "Records are tagged with a controlled vocabulary and linked to "
         "related records in thematic clusters."),
        ("Publish: ", "Records appear on the web dashboard and are available for briefings and search."),
    ]
    for prefix, text in steps:
        add_bullet(doc, text, bold_prefix=prefix)

    # ================================================================
    # WHAT'S IN A RECORD
    # ================================================================
    doc.add_heading("What's in a Record?", level=2)

    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(4)
    p5.add_run(
        "Every record captures the same structured information, regardless of source:"
    ).font.size = Pt(10.5)

    record_fields = [
        ("Title & Date: ", "What happened and when."),
        ("Summary: ", "A 2-4 sentence AI-generated overview focused on what matters for city operations."),
        ("Departments: ", "Which city departments are involved (e.g., Department of Public Works, Syracuse Police Department)."),
        ("People & Organizations: ", "Named individuals, organizations, and city systems mentioned."),
        ("Topic Tags: ", "Standardized categories (e.g., budget, infrastructure, police, lead-remediation)."),
        ("Sentiment: ", "Overall tone — positive, critical, mixed, procedural, etc."),
        ("Political Signals: ", "Detected patterns like championship, opposition, scrutiny, or budget commitment."),
        ("Cluster: ", "Which story or theme this record belongs to (e.g., CLUSTER-BUDGET-BATTLE-2025)."),
    ]
    for prefix, text in record_fields:
        add_bullet(doc, text, bold_prefix=prefix)

    # ================================================================
    # PAGE BREAK — HOW TO USE IT
    # ================================================================
    doc.add_page_break()

    doc.add_heading("How to Use the Context Engine", level=1)

    # -- Dashboard --
    doc.add_heading("The Web Dashboard", level=2)

    p6 = doc.add_paragraph()
    p6.paragraph_format.space_after = Pt(4)
    run_url = p6.add_run("https://conorpmuldoon-cpm.github.io/context-engine/")
    run_url.font.size = Pt(10.5)
    run_url.font.color.rgb = RGBColor(0x0B, 0x5A, 0xB5)
    run_url.underline = True

    p7 = doc.add_paragraph()
    p7.paragraph_format.space_after = Pt(4)
    p7.add_run(
        "The dashboard is the primary way to interact with the Context Engine. "
        "It updates automatically every day. You can:"
    ).font.size = Pt(10.5)

    dash_items = [
        ("Search records ", "by keyword — type a person's name, department, topic, or record ID into the search bar."),
        ("Filter ", "by source type (news, council, press release), department, topic tag, sentiment, or political signal."),
        ("Browse clusters ", "to see groups of related records that tell a story. "
         "For example, the \"Budget Battle\" cluster has 31 records tracking the 2025-2026 budget process."),
        ("Read briefings ", "— AI-generated intelligence packages on specific departments or topics. "
         "These are multi-page documents with executive summaries, council activity, budget context, and open questions."),
        ("View stats ", "on the home page: total records, records by source type, department activity, trending tags."),
    ]
    for prefix, text in dash_items:
        add_bullet(doc, text, bold_prefix=prefix)

    # -- Requesting a Briefing --
    doc.add_heading("Requesting a Briefing", level=2)

    p8 = doc.add_paragraph()
    p8.paragraph_format.space_after = Pt(4)
    p8.add_run(
        "When you need a synthesized intelligence package on a department, topic, or issue:"
    ).font.size = Pt(10.5)

    briefing_steps = [
        'Click "Request a Briefing" on the dashboard (this opens a form on GitHub).',
        "Choose the briefing type: department briefing, topic/theme briefing, cluster briefing, or pre-interview package.",
        'Enter the target — for example, "Department of Public Works" or "lead remediation" or a specific cluster name.',
        "Select a time range (last 3 months, 6 months, 12 months, or all records).",
        "Add any context — upcoming meetings, specific questions you want answered, people to focus on.",
        "Submit the form. The system will automatically score all records for relevance, send the top matches "
        "to AI for synthesis, generate the briefing, and publish it to the dashboard — typically within a few minutes.",
    ]
    for i, text in enumerate(briefing_steps):
        add_bullet(doc, f"{text}", bold_prefix=f"{i + 1}. " if False else None)
        # Use numbered style instead
    # Actually, let's redo as numbered
    # Remove the bullets we just added
    for _ in range(len(briefing_steps)):
        doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)

    for i, text in enumerate(briefing_steps):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"

    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_after = Pt(4)
    run_note = p_note.add_run(
        "Briefings include an executive summary, recent council activity, news coverage, "
        "budget context, active initiatives, a political signals table, and open questions — "
        "all with specific record ID citations you can look up on the dashboard."
    )
    run_note.font.size = Pt(10.5)
    run_note.italic = True

    # -- Submitting a Record --
    doc.add_heading("Submitting a Record Manually", level=2)

    p9 = doc.add_paragraph()
    p9.paragraph_format.space_after = Pt(4)
    p9.add_run(
        "If you encounter information the automated collectors haven't picked up — "
        "a meeting note, a budget observation, a relevant article from another source — "
        "you can add it manually:"
    ).font.size = Pt(10.5)

    add_bullet(doc, 'Click "Submit a Record" on the dashboard.', bold_prefix=None)
    add_bullet(doc, "Fill in the title, source URL, source type, publication date, and a 2-4 sentence summary.", bold_prefix=None)
    add_bullet(doc, "Optionally add department names and topic tags.", bold_prefix=None)
    add_bullet(doc, "Submit. The record will be reviewed and added to the store.", bold_prefix=None)

    # -- Weekly Digest --
    doc.add_heading("Weekly Digest Email", level=2)

    p10 = doc.add_paragraph()
    p10.paragraph_format.space_after = Pt(4)
    p10.add_run(
        "Every Monday morning, the system sends an AI-synthesized email digest covering the past week's "
        "new records. It includes an at-a-glance summary, key takeaways with record citations, "
        "emerging issues, cluster updates, and items to watch. No action needed — it arrives automatically."
    ).font.size = Pt(10.5)

    # -- Tips --
    doc.add_heading("Tips", level=2)

    tips = [
        ("Search by record ID ", "(e.g., CTX-COUNCIL-2026-00056) to jump directly to a specific record."),
        ("Use cluster names ", "to explore a full story arc — clusters group records from different sources "
         "that relate to the same issue."),
        ("Request briefings before interviews ", "— the pre-interview package option is designed for "
         "Discovery engagements and will focus on the specific department and its recent activity."),
        ("Flag useful or noisy records ", "— your feedback helps calibrate future relevance scoring."),
    ]
    for prefix, text in tips:
        add_bullet(doc, text, bold_prefix=prefix)

    # ================================================================
    # SAVE
    # ================================================================
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
